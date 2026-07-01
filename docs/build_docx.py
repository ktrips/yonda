#!/usr/bin/env python3
"""
cursor_claude_amazon_app.md → Kindle判型 (210×257mm) docx 変換スクリプト
・表紙・自動目次（ハイパーリンク付き）・フッターページ番号付き
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ──────────────────────────────────────────────
# ページ設定
# ──────────────────────────────────────────────

def set_page_layout(section, width_mm=210, height_mm=257,
                    top=15, bottom=15, inner=18, outer=12):
    section.page_width  = Mm(width_mm)
    section.page_height = Mm(height_mm)
    section.top_margin    = Mm(top)
    section.bottom_margin = Mm(bottom)
    section.left_margin   = Mm(inner)
    section.right_margin  = Mm(outer)


# ──────────────────────────────────────────────
# スタイル設定
# ──────────────────────────────────────────────

def setup_styles(doc: Document):
    styles = doc.styles

    # Normal
    n = styles["Normal"]
    n.font.name = "Hiragino Kaku Gothic ProN"
    n.font.size = Pt(10)
    n.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    n.paragraph_format.space_after  = Pt(5)
    n.paragraph_format.line_spacing = Pt(16)

    # H1
    h1 = styles["Heading 1"]
    h1.font.name  = "Hiragino Kaku Gothic ProN"
    h1.font.size  = Pt(17)
    h1.font.bold  = True
    h1.font.color.rgb = RGBColor(0x16, 0x3A, 0x7A)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True   # H1 は常に新ページから

    # H2
    h2 = styles["Heading 2"]
    h2.font.name  = "Hiragino Kaku Gothic ProN"
    h2.font.size  = Pt(13)
    h2.font.bold  = True
    h2.font.color.rgb = RGBColor(0x2A, 0x5A, 0xA0)
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after  = Pt(5)
    h2.paragraph_format.keep_with_next = True

    # H3
    h3 = styles["Heading 3"]
    h3.font.name  = "Hiragino Kaku Gothic ProN"
    h3.font.size  = Pt(11)
    h3.font.bold  = True
    h3.font.color.rgb = RGBColor(0x35, 0x65, 0xAA)
    h3.paragraph_format.space_before = Pt(11)
    h3.paragraph_format.space_after  = Pt(3)
    h3.paragraph_format.keep_with_next = True

    # H4
    h4 = styles["Heading 4"]
    h4.font.name   = "Hiragino Kaku Gothic ProN"
    h4.font.size   = Pt(10.5)
    h4.font.bold   = True
    h4.font.italic = False
    h4.font.color.rgb = RGBColor(0x44, 0x44, 0x55)
    h4.paragraph_format.space_before = Pt(9)
    h4.paragraph_format.space_after  = Pt(2)
    h4.paragraph_format.keep_with_next = True

    # Code Block
    _ensure_style(styles, "Code Block", WD_STYLE_TYPE.PARAGRAPH, "Normal",
                  font_name="Courier New", font_size=7.5,
                  color=RGBColor(0x1E, 0x1E, 0x2E),
                  space_before=0, space_after=0,
                  left_indent=Mm(3), line_spacing=Pt(10.5))

    # Block Quote
    _ensure_style(styles, "Block Quote", WD_STYLE_TYPE.PARAGRAPH, "Normal",
                  font_name="Hiragino Kaku Gothic ProN", font_size=9.5,
                  color=RGBColor(0x55, 0x55, 0x66), italic=True,
                  space_before=6, space_after=6,
                  left_indent=Mm(6))

    # List Bullet
    _ensure_style(styles, "List Bullet Yonda", WD_STYLE_TYPE.PARAGRAPH, "Normal",
                  font_name="Hiragino Kaku Gothic ProN", font_size=10,
                  space_before=1, space_after=2,
                  left_indent=Mm(6))


def _ensure_style(styles, name, stype, base, font_name=None, font_size=None,
                  color=None, italic=False, bold=False,
                  space_before=None, space_after=None,
                  left_indent=None, line_spacing=None):
    if name in [s.name for s in styles]:
        return
    s = styles.add_style(name, stype)
    s.base_style = styles[base]
    if font_name:   s.font.name = font_name
    if font_size:   s.font.size = Pt(font_size)
    if color:       s.font.color.rgb = color
    if italic:      s.font.italic = True
    if bold:        s.font.bold = True
    pf = s.paragraph_format
    if space_before is not None: pf.space_before = Pt(space_before)
    if space_after  is not None: pf.space_after  = Pt(space_after)
    if left_indent  is not None: pf.left_indent  = left_indent
    if line_spacing is not None: pf.line_spacing = line_spacing


# ──────────────────────────────────────────────
# フッター：ページ番号
# ──────────────────────────────────────────────

def add_page_numbers_to_footer(section, start_num: int = 1):
    """フッター中央にページ番号フィールドを追加"""
    footer = section.footer
    footer.is_linked_to_previous = False

    # 既存段落をクリア
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ページ番号フィールド
    run = fp.add_run()
    _add_fld_char(run, ' PAGE ')
    run.font.name = "Hiragino Kaku Gothic ProN"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x88)

    # ページ番号の開始値を設定
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), str(start_num))
    sectPr.append(pgNumType)


def _add_fld_char(run, instr: str):
    """シンプルなフィールド文字列を run に追加"""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrEl = OxmlElement("w:instrText")
    instrEl.set(qn("xml:space"), "preserve")
    instrEl.text = instr
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrEl)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _blank_footer(section):
    """カバーページのフッターを空白にする"""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()


# ──────────────────────────────────────────────
# 表紙ページ
# ──────────────────────────────────────────────

def add_cover_page(doc: Document):
    """表紙：タイトル・サブタイトル・日付"""
    section = doc.sections[0]
    set_page_layout(section)
    section.different_first_page_header_footer = True
    _blank_footer(section)   # 表紙はフッターなし

    def _center_para(text, size, bold=False, color=None, space_before=0, space_after=8):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name = "Hiragino Kaku Gothic ProN"
        r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = color
        return p

    # 上部余白
    for _ in range(6):
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(0)

    # メインタイトル
    _center_para(
        "Cursor + Claude で\nAmazon 連携読書記録アプリを作る",
        24, bold=True,
        color=RGBColor(0x16, 0x3A, 0x7A),
        space_before=0, space_after=12
    )

    # サブタイトル
    _center_para(
        "実践ガイド",
        18, bold=True,
        color=RGBColor(0x2A, 0x5A, 0xA0),
        space_before=0, space_after=20
    )

    # 区切り線
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(6)
    sep.paragraph_format.space_after  = Pt(6)
    r = sep.add_run("─" * 28)
    r.font.color.rgb = RGBColor(0xBB, 0xCC, 0xDD)
    r.font.size = Pt(10)

    # キャッチコピー
    _center_para(
        "Audible・Kindle・公共図書館・紙の本を一元管理し\nAI が書評と選書を支援する Web アプリを\nAI エディタで実際に作る方法を解説する実践書",
        11,
        color=RGBColor(0x44, 0x55, 0x66),
        space_before=8, space_after=40
    )

    # 下部余白を埋めてから日付
    _center_para(
        "2026年6月",
        10,
        color=RGBColor(0x88, 0x88, 0x99),
        space_before=60, space_after=4
    )
    _center_para(
        "yonda プロジェクト",
        10, bold=True,
        color=RGBColor(0x55, 0x66, 0x88),
        space_before=0, space_after=4
    )

    # 改ページ（目次へ）
    doc.add_page_break()


# ──────────────────────────────────────────────
# 目次（Word TOC フィールド + Kindle ハイパーリンク対応）
# ──────────────────────────────────────────────

def add_toc_page(doc: Document):
    """
    Word の自動TOCフィールドを挿入。
    \\h  = ハイパーリンク付き（Kindleの章ジャンプに対応）
    \\z  = Webレイアウトでは点線タブを非表示
    \\u  = アウトラインレベルを使用
    \\o "1-3" = 見出し1〜3を収録
    """
    # 「目次」見出し
    p_title = doc.add_paragraph("目次")
    p_title.style = "Normal"
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(16)
    r = p_title.runs[0]
    r.font.name  = "Hiragino Kaku Gothic ProN"
    r.font.size  = Pt(16)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x16, 0x3A, 0x7A)

    # TOC フィールド挿入
    p_toc = doc.add_paragraph()
    p_toc.style = "Normal"
    run = p_toc.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")

    # プレースホルダー（Wordで開くと自動更新される）
    placeholder = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    noProof = OxmlElement("w:noProof")
    rPr.append(noProof)
    placeholder.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = "（Wordで開いてフィールドを更新してください / Press Ctrl+A → F9 to update）"
    placeholder.append(t)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instr)
    run._r.append(fldChar_sep)
    run._r.append(placeholder)
    run._r.append(fldChar_end)

    # 目次の後で改ページ
    doc.add_page_break()


# ──────────────────────────────────────────────
# コードブロック装飾
# ──────────────────────────────────────────────

def _add_code_shading(para, fill="F3F4F6"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _add_border_bottom(para, color="CCCCCC", sz="4"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────────────────────────────────────────────
# インラインマークアップ
# ──────────────────────────────────────────────

def parse_inline(para, text: str):
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`|_(.+?)_)')
    pos = 0
    for m in pattern.finditer(text):
        start, end = m.start(), m.end()
        if start > pos:
            r = para.add_run(text[pos:start])
            r.font.name = "Hiragino Kaku Gothic ProN"
        full = m.group(0)
        if full.startswith("**"):
            r = para.add_run(m.group(2))
            r.bold = True
            r.font.name = "Hiragino Kaku Gothic ProN"
        elif full.startswith("`"):
            r = para.add_run(m.group(3))
            r.font.name  = "Courier New"
            r.font.size  = Pt(8.5)
            r.font.color.rgb = RGBColor(0xBF, 0x1F, 0x4A)
        else:
            r = para.add_run(m.group(4))
            r.italic = True
            r.font.name = "Hiragino Kaku Gothic ProN"
        pos = end
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.font.name = "Hiragino Kaku Gothic ProN"


# ──────────────────────────────────────────────
# Markdownテーブル
# ──────────────────────────────────────────────

def build_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        if re.match(r'^\|[-| :]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return

    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"

    col_w = Mm(170 / ncols)
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_w

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri]
        for ci in range(min(len(row_data), ncols)):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            if ri == 0:
                r = p.add_run(row_data[ci])
                r.bold = True
                r.font.name = "Hiragino Kaku Gothic ProN"
                r.font.size = Pt(8.5)
                # ヘッダー行を薄い青で塗る
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D6E4F7")
                tcPr.append(shd)
            else:
                p.clear()
                parse_inline(p, row_data[ci])
                for r in p.runs:
                    r.font.size = Pt(8.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ──────────────────────────────────────────────
# Markdown → docx 変換本体
# ──────────────────────────────────────────────

def convert_md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    setup_styles(doc)

    # ── 表紙（セクション1）──
    add_cover_page(doc)

    # ── セクション2（目次 + 本文）を追加 ──
    # 表紙の改ページ後に新しいセクションを作る
    # python-docx では add_section() で新セクションを追加できる
    from docx.enum.section import WD_SECTION
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_layout(new_section)
    new_section.different_first_page_header_footer = False
    new_section.footer.is_linked_to_previous = False

    # ── 目次 ──
    add_toc_page(doc)

    # ── 本文フッター（ページ番号）はセクション2に設定 ──
    add_page_numbers_to_footer(new_section, start_num=1)

    # ── Markdown パース ──
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_lines: list[str] = []

    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}

    while i < len(lines):
        line = lines[i]

        # ── コードブロック ──
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                in_code = False
                if code_lang:
                    lp = doc.add_paragraph(f"[{code_lang}]")
                    lp.style = "Normal"
                    for r in lp.runs:
                        r.font.size = Pt(7)
                        r.font.color.rgb = RGBColor(0x99, 0x99, 0xAA)
                        r.font.name = "Courier New"
                    lp.paragraph_format.space_after = Pt(0)
                for j, cl in enumerate(code_lines):
                    cp = doc.add_paragraph(cl if cl else " ")
                    cp.style = "Code Block"
                    _add_code_shading(cp)
                    cp.paragraph_format.space_before = Pt(4 if j == 0 else 0)
                    cp.paragraph_format.space_after  = Pt(6 if j == len(code_lines)-1 else 0)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── テーブル ──
        if line.startswith("|"):
            table_lines.append(line)
            i += 1
            if i >= len(lines) or not lines[i].startswith("|"):
                build_table(doc, table_lines)
                table_lines = []
            continue

        # ── 見出し ──
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', m.group(2).strip())
            doc.add_paragraph(text, style=style_map.get(level, "Heading 4"))
            i += 1
            continue

        # ── 水平線 ──
        if re.match(r'^---+\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _add_border_bottom(p)
            i += 1
            continue

        # ── 引用 ──
        if line.startswith("> "):
            p = doc.add_paragraph(style="Block Quote")
            parse_inline(p, line[2:].strip())
            i += 1
            continue

        # ── 箇条書き ──
        m_li = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if m_li:
            indent = len(m_li.group(1))
            is_num = bool(re.match(r'\d+\.', m_li.group(2)))
            prefix = m_li.group(2) + " " if is_num else "• "
            p = doc.add_paragraph(style="List Bullet Yonda")
            p.paragraph_format.left_indent = Mm(6 + indent * 3)
            r0 = p.add_run(prefix)
            r0.font.name = "Hiragino Kaku Gothic ProN"
            parse_inline(p, m_li.group(3))
            i += 1
            continue

        # ── 空行 ──
        if not line.strip():
            i += 1
            continue

        # ── 通常段落 ──
        p = doc.add_paragraph(style="Normal")
        parse_inline(p, line.strip())
        i += 1

    doc.save(str(docx_path))
    print(f"✅ 保存完了: {docx_path}")
    size_kb = docx_path.stat().st_size // 1024
    print(f"   サイズ: {size_kb} KB")
    print(f"   段落数: {len(doc.paragraphs)}")
    print()
    print("📌 次のステップ:")
    print("   1. Word で開く")
    print("   2. Ctrl+A（全選択）→ F9 で目次フィールドを更新")
    print("   3. 目次のページ番号が反映される")
    print("   4. Kindle Direct Publishing にアップロード可能")


if __name__ == "__main__":
    root      = Path(__file__).parent
    md_path   = root / "cursor_claude_amazon_app.md"
    docx_path = root / "cursor_claude_amazon_app.docx"
    convert_md_to_docx(md_path, docx_path)
